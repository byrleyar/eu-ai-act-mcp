from docx import Document
import zipfile
import io


def load_docx_from_template(path_or_file) -> Document:
    """
    Loads a .docx file.
    """
    return Document(path_or_file)

def process_tables(tables, display_data):
    """
    Recursively process tables to handle nested tables.
    """
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs in the cell
                for paragraph in cell.paragraphs:
                     for key, value in display_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                
                # Recursively process nested tables in the cell
                if cell.tables:
                    process_tables(cell.tables, display_data)

def fill_template(template_path, output_path, data: dict):
    """
    Fills a docx template by replacing {{key}} with value from data.
    template_path and output_path can be file paths (str) or file-like objects.
    """
    # Load template safely
    doc = load_docx_from_template(template_path)
    
    # Prepare visually friendly data map (converting yes/no to checkboxes)
    display_data = {}
    for k, v in data.items():
        val_str = str(v).strip().lower()
        if val_str == "yes":
            display_data[k] = "☑" # Checked box
        elif val_str == "no":
            display_data[k] = "☐" # Unchecked box
        else:
            display_data[k] = str(v)
    
    # Process body paragraphs
    for paragraph in doc.paragraphs:
        for key, value in display_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Process tables recursively
    process_tables(doc.tables, display_data)

    doc.save(output_path)


